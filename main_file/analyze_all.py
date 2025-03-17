
import ssl

try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    # Legacy Python that doesn't verify HTTPS certificates by default
    # Старая версия Python, которая по умолчанию не проверяет HTTPS-сертификаты
    pass
else:
    # Handle target environment that doesn't support HTTPS verification
    # Обработка целевого окружения, которое не поддерживает проверку HTTPS
    ssl._create_default_https_context = _create_unverified_https_context

import re
import spacy
import os
from docx import Document
import PyPDF2
from bs4 import BeautifulSoup
import traceback
import logging
from sumy.parsers.plaintext import PlaintextParser
from sumy.summarizers.text_rank import TextRankSummarizer
from langdetect import detect
import jieba
import nltk
from nltk.stem import WordNetLemmatizer, PorterStemmer

nltk.download('wordnet')


# 定义不同语言的停用词
# Определение стоп-слов для разных языков
# 这里假设已经有中文停用词文件 stopwords_zh.txt
# Здесь предполагается, что уже есть файл с китайскими стоп-словами stopwords_zh.txt
def load_stopwords(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return set([line.strip() for line in f.readlines()])


try:
    stopwords_zh = load_stopwords('stopwords_zh.txt')
except FileNotFoundError:
    stopwords_zh = set()

STOPWORDS = {
    'zh': stopwords_zh,
    'en': set(['a', 'an', 'the', 'in', 'on', 'at', 'for', 'to', 'of', 'and']),
    'ru': set(['и', 'в', 'на', 'за', 'с', 'о', 'от', 'для'])
}

# 配置日志记录，设置日志级别为INFO，定义日志格式，包括时间、日志级别和消息内容
# Конфигурация логирования, установка уровня логирования INFO, определение формата лога, который включает время, уровень логирования и содержание сообщения
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


# 解析.pdf文件的函数，接受文件路径作为参数
# Функция для разбора PDF-файла, принимает путь к файлу в качестве параметра
def parse_pdf(file_path):
    try:
        text = ""
        # 以二进制只读模式打开PDF文件
        # Открыть PDF-файл в бинарном режиме для чтения
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            # 遍历PDF文件的每一页
            # Пробежаться по каждой странице PDF-файла
            for page in reader.pages:
                # 提取每一页的文本内容并累加到text变量中
                # Извлечь текст каждой страницы и добавить его в переменную text
                text += page.extract_text()
        return text
    # 如果在解析过程中出现异常，记录错误日志并返回空字符串
    # Если во время разбора возникает исключение, записать ошибку в журнал и вернуть пустую строку
    except Exception as e:
        logging.error(f"解析.pdf文件时出错: {e}")
        return ""


# 解析.docx和部分.doc文件的函数，接受文件路径作为参数
# Функция для разбора DOCX и некоторых DOC-файлов, принимает путь к файлу в качестве параметра
def parse_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        # 遍历.docx文件中的每一个段落
        # Пробежаться по каждому абзацу в DOCX-файле
        for para in doc.paragraphs:
            # 将每个段落的文本添加到full_text列表中
            # Добавить текст каждого абзаца в список full_text
            full_text.append(para.text)
        # 使用换行符连接列表中的所有文本并返回
        # Соединить все тексты в списке с помощью символа переноса строки и вернуть результат
        return "\n".join(full_text)
    # 如果在解析过程中出现异常，记录错误日志并返回空字符串
    # Если во время разбора возникает исключение, записать ошибку в журнал и вернуть пустую строку
    except Exception as e:
        logging.error(f"解析.docx文件时出错: {e}")
        return ""


# 解析.html文件的函数，接受文件路径作为参数
# Функция для разбора HTML-файла, принимает путь к файлу в качестве параметра
def parse_html(file_path):
    try:
        # 以只读模式打开HTML文件，指定编码为utf-8
        # Открыть HTML-файл в режиме для чтения, указать кодировку UTF-8
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
            # 使用BeautifulSoup解析HTML内容，指定解析器为html.parser
            # Разобрать содержимое HTML с использованием BeautifulSoup, указать парсер html.parser
            soup = BeautifulSoup(html_content, 'html.parser')
            # 提取文本内容并返回
            # Извлечь текстовое содержимое и вернуть его
            text_content = soup.get_text()
            return text_content
    # 如果文件未找到，记录错误日志并返回空字符串
    # Если файл не найден, записать ошибку в журнал и вернуть пустую строку
    except FileNotFoundError:
        logging.error(f"文件未找到: {file_path}")
        return ""
    # 如果在解析过程中出现其他异常，记录错误日志并返回空字符串
    # Если во время разбора возникает другое исключение, записать ошибку в журнал и вернуть пустую строку
    except Exception as e:
        logging.error(f"解析HTML文件时出错: {e}")
        return ""


# 尝试解析.djvu文件的函数（需要安装pydjvu库），接受文件路径作为参数
# Функция для попытки разбора DJVU-файла (требуется установка библиотеки pydjvu), принимает путь к файлу в качестве параметра
def parse_djvu(file_path):
    try:
        from pydjvu.lib import DjVuDocument
        doc = DjVuDocument(file_path)
        text = ""
        # 遍历.djvu文件的每一页
        # Пробежаться по каждой странице DJVU-файла
        for page in doc:
            # 提取每一页的文本内容并累加到text变量中
            # Извлечь текст каждой страницы и добавить его в переменную text
            text += page.get_text()
        return text
    # 如果未安装pydjvu库，捕获ImportError异常，记录错误日志并返回空字符串
    # Если библиотека pydjvu не установлена, перехватить исключение ImportError, записать ошибку в журнал и вернуть пустую строку
    except ImportError:
        logging.error("未安装pydjvu库，无法解析.djvu文件。")
        return ""
    # 如果在解析过程中出现其他异常，记录错误日志并返回空字符串
    # Если во время разбора возникает другое исключение, записать ошибку в журнал и вернуть пустую строку
    except Exception as e:
        logging.error(f"解析.djvu文件时出错: {e}")
        return ""


# 使用NER（命名实体识别）识别品牌名称、热门关键词、流行趋势的函数，接受文本作为参数
# Функция для определения имен брендов, популярных ключевых слов и трендов с использованием NER (определение именованных сущностей), принимает текст в качестве параметра
def ner_analysis(text):
    try:
        # 使用langdetect库检测文本的语言
        # Определить язык текста с использованием библиотеки langdetect
        lang = detect(text)
        # 根据检测到的语言加载相应的spacy模型
        # Загрузить соответствующую модель spacy в зависимости от определенного языка
        if lang == 'zh':
            nlp = spacy.load("zh_core_web_sm")
        elif lang == 'ru':
            nlp = spacy.load("ru_core_news_sm")
        else:
            nlp = spacy.load("en_core_web_sm")
        doc = nlp(text)
        brand_names = []
        keywords = []
        # 遍历文档中的每一个实体
        # Пробежаться по каждой сущности в документе
        for ent in doc.ents:
            # 如果实体的标签是"ORG"（组织）或"PRODUCT"（产品），将其添加到品牌名称列表中
            # Если метка сущности - "ORG" (организация) или "PRODUCT" (продукт), добавить ее в список имен брендов
            if ent.label_ in ["ORG", "PRODUCT"]:
                brand_names.append(ent.text)
        # 调用extract_keywords函数提取关键词
        # Вызвать функцию extract_keywords для извлечения ключевых слов
        keywords = extract_keywords(text)
        return brand_names, keywords
    # 如果在NER分析过程中出现异常，记录错误日志并返回空的品牌名称和关键词列表
    # Если во время анализа NER возникает исключение, записать ошибку в журнал и вернуть пустые списки имен брендов и ключевых слов
    except Exception as e:
        logging.error(f"NER分析时出错: {e}")
        return [], []


# 使用TextRank算法提取关键词的函数，接受文本作为参数
# Функция для извлечения ключевых слов с использованием алгоритма TextRank, принимает текст в качестве параметра
def extract_keywords(text):
    try:
        # 使用langdetect库检测文本的语言
        # Определить язык текста с использованием библиотеки langdetect
        lang = detect(text)
        if lang == 'zh' or lang == 'zh-cn':
            # 中文分词
            # Разделение на токены для китайского языка
            words = jieba.lcut(text)
            new_text = " ".join(words)
        else:
            # 英文使用空格分割
            # Разделение на токены для английского языка с использованием пробела
            new_text = text

        # 手动创建分词器，实现to_sentences和to_words方法
        # Создать токенизатор вручную, реализовать методы to_sentences и to_words
        class SimpleTokenizer:
            def to_sentences(self, text):
                # 简单地按句号分割句子
                # Простое разделение предложений по точке
                return text.split('.')

            def to_words(self, text):
                return text.split()

        tokenizer = SimpleTokenizer()
        # 创建解析器，传入文本和分词器
        # Создать парсер, передать в него текст и токенизатор
        parser = PlaintextParser.from_string(new_text, tokenizer)
        summarizer = TextRankSummarizer()
        # 使用TextRankSummarizer提取10个关键句子
        # Извлечь 10 ключевых предложений с использованием TextRankSummarizer
        sentences = summarizer(parser.document, 10)
        keywords = []
        # 遍历提取的每一个句子
        # Пробежаться по каждому извлеченному предложению
        for sentence in sentences:
            # 遍历句子中的每一个单词
            # Пробежаться по каждому слову в предложении
            for word in sentence.words:
                # 过滤停用词
                # Фильтрация стоп-слов
                if word not in STOPWORDS.get(lang, set()) and word not in keywords:
                    keywords.append(word)
        return keywords
    # 如果在关键词提取过程中出现异常，记录错误日志
    # Если во время извлечения ключевых слов возникает исключение, записать ошибку в журнал
    except Exception as e:
        logging.error(f"关键词提取时出错: {e}")
        return []


# 句法分析函数，接受文本作为参数
# Функция для синтаксического анализа, принимает текст в качестве параметра
def syntax_analysis(text):
    try:
        # 使用langdetect库检测文本的语言
        # Определить язык текста с использованием библиотеки langdetect
        lang = detect(text)
        # 根据检测到的语言加载相应的spacy模型
        # Загрузить соответствующую модель spacy в зависимости от определенного языка
        if lang == 'zh':
            nlp = spacy.load("zh_core_web_sm")
        elif lang == 'ru':
            nlp = spacy.load("ru_core_news_sm")
        else:
            nlp = spacy.load("en_core_web_sm")
        doc = nlp(text)
        # 遍历文档中的每一个词元（token）
        # Пробежаться по каждому токену в документе
        for token in doc:
            # 打印词元的文本、词性、依存关系和头词
            # Вывести текст токена, часть речи, зависимость и головное слово
            print(f"词: {token.text}, 词性: {token.pos_}, 依存关系: {token.dep_}, 头词: {token.head.text}")
    # 如果在句法分析过程中出现异常，记录错误日志
    # Если во время синтаксического анализа возникает исключение, записать ошибку в журнал
    except Exception as e:
        logging.error(f"句法分析时出错: {e}")


# 词法化函数
# Функция для лемматизации
def lemmatize_text(text):
    lemmatizer = WordNetLemmatizer()
    words = text.split()
    lemmatized_words = [lemmatizer.lemmatize(word) for word in words]
    return " ".join(lemmatized_words)


# 词干化函数
# Функция для стемминга
def stem_text(text):
    stemmer = PorterStemmer()
    words = text.split()
    stemmed_words = [stemmer.stem(word) for word in words]
    return " ".join(stemmed_words)


# 主函数，程序入口
# Главная функция, точка входа в программу
def main():
    file_path = input("请输入文件路径: ")
    # 如果文件路径为空或文件不存在，记录错误日志并返回
    # Если путь к файлу пустой или файл не существует, записать ошибку в журнал и вернуться
    if not file_path or not os.path.exists(file_path):
        logging.error("输入的文件路径无效。")
        return

    file_extension = file_path.split('.')[-1].lower()
    # 根据文件扩展名调用相应的解析函数
    # Вызвать соответствующую функцию разбора в зависимости от расширения файла
    if file_extension == "pdf":
        text = parse_pdf(file_path)
    elif file_extension == "docx" or file_extension == "doc":
        text = parse_docx(file_path)
    elif file_extension == "html":
        text = parse_html(file_path)
    elif file_extension == "djvu":
        text = parse_djvu(file_path)
    else:
        logging.error("不支持的文件类型。")
        return

    # 预处理文本，去除多余空格和空字符
    # Предварительная обработка текста, удаление лишних пробелов и пустых символов
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)

    if text:
        # 词法化和词干化
        # Лемматизация и стемминг
        lemmatized_text = lemmatize_text(text)
        stemmed_text = stem_text(text)

        brand_names, keywords = ner_analysis(text)
        print("识别到的品牌名称:", brand_names)
        print("识别到的热门关键词:", keywords)
        syntax_analysis(text)


if __name__ == "__main__":
    try:
        main()
    # 如果在主函数执行过程中出现异常，记录错误日志并打印异常堆栈信息
    # Если во время выполнения главной функции возникает исключение, записать ошибку в журнал и вывести стек_trace ошибки
    except Exception as e:
        logging.error(f"发生错误: {e}")
        traceback.print_exc()